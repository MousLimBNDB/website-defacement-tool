import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_formatted_text(paragraph, text, base_font_size=11, font_name="Segoe UI"):
    # Parse bold (**text**) and code (`text`) inside text string
    pattern = re.compile(r'(\*\*.+?\*\*|`.+?`)')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(base_font_size)
        
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(base_font_size - 0.5)
            run.font.color.rgb = RGBColor(0xC5, 0x30, 0x30)
        else:
            run.text = token
            run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Set page margins (1 inch = 1440 dxa)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    base_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_lines = []

    for line in lines:
        stripped = line.strip()

        # Handle Code Block fences
        if stripped.startswith("```"):
            if in_code_block:
                # Close code block
                in_code_block = False
                code_text = "\n".join(code_block_lines)
                code_block_lines = []

                # Add a shaded callout table for code block
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                cell.width = Inches(6.5)
                set_cell_background(cell, "F7FAFC")
                set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            else:
                in_code_block = True
                code_block_lines = []
            continue

        if in_code_block:
            code_block_lines.append(line.rstrip("\r\n"))
            continue

        if not stripped:
            continue

        # Horizontal Rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E0"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)
            continue

        # Document Title (H1)
        if stripped.startswith("# "):
            title_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(title_text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x29, 0x4A)
            continue

        # Heading 2 (##)
        if stripped.startswith("## "):
            h_text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            continue

        # Heading 3 (###)
        if stripped.startswith("### "):
            h_text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
            continue

        # Check for Image tags: ![alt](src)
        img_match = re.match(r'^\s*!\[(.*?)\]\((.*?)\)\s*$', stripped)
        if img_match:
            alt_text, img_src = img_match.group(1), img_match.group(2)
            resolved_img_path = os.path.normpath(os.path.join(base_dir, img_src))
            
            if os.path.exists(resolved_img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                
                run = p.add_run()
                run.add_picture(resolved_img_path, width=Inches(6.2))
                
                # Image Caption
                if alt_text:
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_after = Pt(12)
                    cap_run = cap_p.add_run(f"Figure: {alt_text}")
                    cap_run.font.name = "Segoe UI"
                    cap_run.font.size = Pt(9.5)
                    cap_run.font.italic = True
                    cap_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
            else:
                print(f"Warning: Image file not found: {resolved_img_path}")
            continue

        # Bullet List Items
        if stripped.startswith("- ") or stripped.startswith("* "):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, item_text)
            continue

        # Sub-bullet items (e.g. 2 spaces or 4 spaces + dash)
        if re.match(r'^\s+[\-\*]\s+', line):
            item_text = re.sub(r'^\s+[\-\*]\s+', '', line).strip()
            p = doc.add_paragraph(style='List Bullet 2' if 'List Bullet 2' in doc.styles else 'List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, item_text)
            continue

        # Numbered List Items
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            item_text = num_match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, item_text)
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped)

    doc.save(docx_path)
    print(f"Successfully generated DOCX document: {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx(
        md_path="RepportIterations/report_v1.md",
        docx_path="RepportIterations/report_v1.docx"
    )
